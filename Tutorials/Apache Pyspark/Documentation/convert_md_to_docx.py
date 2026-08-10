from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

md_path = Path(r'c:\Mukesh\Self-Practice\DE Upskill Practice\Sessions\Session 8 - Tasks\Apache PySpark\PySpark Notebook Documentation.md')
out_path = Path(r'c:\Mukesh\Self-Practice\DE Upskill Practice\Sessions\Session 8 - Tasks\Apache PySpark\PySpark Notebook Documentation.docx')

if not md_path.exists():
    raise FileNotFoundError(f'Markdown file not found: {md_path}')

lines = md_path.read_text(encoding='utf-8').splitlines()
doc = Document()

in_code_block = False
for raw_line in lines:
    line = raw_line.rstrip()
    if line.startswith('```'):
        in_code_block = not in_code_block
        continue

    if in_code_block:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        continue

    if not line.strip():
        doc.add_paragraph()
        continue

    if line.startswith('# '):
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        run.bold = True
        run.font.size = Pt(22)
    elif line.startswith('## '):
        p = doc.add_paragraph()
        run = p.add_run(line[3:])
        run.bold = True
        run.font.size = Pt(16)
    elif line.startswith('### '):
        p = doc.add_paragraph()
        run = p.add_run(line[4:])
        run.bold = True
        run.font.size = Pt(13)
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(line[2:])
    elif line.startswith('> '):
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        run.italic = True
    else:
        p = doc.add_paragraph()
        p.add_run(line)

doc.save(out_path)
print(f'Created: {out_path}')
